# -*- coding: utf-8 -*-
"""Minimal Markdown -> Word (.docx) converter for SparkOrbit docs."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="宋体", size=12, bold=False, code=False):
    run.bold = bold
    run.font.size = Pt(size)
    if code:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    else:
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_inline(paragraph, text, base_size=12):
    """Parse **bold**, `code`, plain text into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|[^*`]+)")
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size - 1, code=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def parse_table_row(line: str):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    cells = parse_table_row(line)
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def convert(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    i = 0
    in_code = False
    code_lang = ""
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # fenced code
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_buf = []
            else:
                in_code = False
                block = "\n".join(code_buf)
                label = f"[代码块: {code_lang}]" if code_lang else "[代码块]"
                if code_lang.lower() == "mermaid":
                    p = doc.add_paragraph()
                    run = p.add_run("[Mermaid 图 — 请在 Markdown 源文件中查看渲染效果]")
                    set_run_font(run, size=10, bold=True)
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(block[:2000] + ("…" if len(block) > 2000 else ""))
                    set_run_font(run2, size=9, code=True)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(label)
                    set_run_font(run, size=9, bold=True)
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(block if block else "(空)")
                    set_run_font(run2, size=9, code=True)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # skip horizontal rule
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", line.strip()):
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            # strip trailing markdown links artifacts
            title = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", title)
            if level == 1:
                p = doc.add_heading("", level=0)
                run = p.add_run(title)
                set_run_font(run, name="黑体", size=18, bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_heading("", level=min(level, 4))
                run = p.add_run(title)
                sizes = {2: 16, 3: 14, 4: 12, 5: 12, 6: 11}
                set_run_font(run, name="黑体", size=sizes.get(level, 12), bold=True)
            i += 1
            continue

        # table
        if "|" in line and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            headers = parse_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                if is_table_sep(lines[i]):
                    i += 1
                    continue
                rows.append(parse_table_row(lines[i]))
                i += 1
            cols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline(p, h, base_size=10)
                for run in p.runs:
                    run.bold = True
            for r_idx, row in enumerate(rows):
                for j in range(cols):
                    val = row[j] if j < len(row) else ""
                    cell = table.rows[r_idx + 1].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, val, base_size=10)
            doc.add_paragraph()
            continue

        # unordered list
        m = re.match(r"^(\s*)([-*+])\s+(.*)$", line)
        if m:
            content = m.group(3)
            content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content, base_size=12)
            i += 1
            continue

        # ordered list
        m = re.match(r"^(\s*)(\d+)[.)]\s+(.*)$", line)
        if m:
            content = m.group(3)
            content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, content, base_size=12)
            i += 1
            continue

        # blockquote
        if line.startswith(">"):
            content = line.lstrip("> ").strip()
            content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_run_font(run, size=11)
            run.italic = True
            i += 1
            continue

        # blank
        if not line.strip():
            i += 1
            continue

        # normal paragraph
        content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line.strip())
        p = doc.add_paragraph()
        add_inline(p, content, base_size=12)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"OK: {out_path.name}")


def main():
    root = Path(__file__).resolve().parents[1] / "docs" / "software-eng"
    names = [
        "SparkOrbit-D1-模块开发卷宗.md",
        "SparkOrbit-D2-用户手册.md",
        "SparkOrbit-D3-操作手册.md",
        "SparkOrbit-E1-测试计划.md",
        "SparkOrbit-E2-测试分析报告.md",
        "SparkOrbit-F1-开发进度月报.md",
        "SparkOrbit-F2-项目开发总结报告.md",
        "SparkOrbit-G1-软件质量保证计划.md",
        "SparkOrbit-G2-软件配置管理计划.md",
    ]
    out_dir = root.parent / "export_docx"
    for name in names:
        md = root / name
        if not md.exists():
            print(f"MISSING: {name}", file=sys.stderr)
            continue
        convert(md, out_dir / (md.stem + ".docx"))


if __name__ == "__main__":
    main()
