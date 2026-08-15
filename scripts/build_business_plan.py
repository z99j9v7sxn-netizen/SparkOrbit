#!/usr/bin/env python3
"""SparkOrbit 星轨学图 — 商业计划书 Word 文档生成脚本"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── 输出目录 ──
OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "submit")
os.makedirs(OUT_DIR, exist_ok=True)
OUT_PATH = os.path.join(OUT_DIR, "SparkOrbit_商业计划书.docx")

# ── 配色与品牌 ──
DARK_BG = RGBColor(0x0B, 0x12, 0x20)
ACCENT = RGBColor(0x3D, 0xB8, 0xA0)
WARN = RGBColor(0xE8, 0xA8, 0x38)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_TEXT = RGBColor(0x1A, 0x23, 0x32)
GRAY_TEXT = RGBColor(0x6B, 0x7C, 0x93)
LIGHT_BG = RGBColor(0xF4, 0xF7, 0xFB)

# ── 团队信息 ──
TEAM_NAME = "爱拼才慧莹"
SCHOOL = "吉林外国语大学"
MEMBERS = [
    {"name": "吴灏贤", "role": "项目负责人 / 全栈技术负责人", "desc": "全面负责系统架构设计、前后端开发、多智能体编排、多模态集成与云部署"},
    {"name": "姜慧", "role": "产品文档负责人 / UI 设计师", "desc": "负责产品需求文档撰写、用户界面视觉设计、软件工程文档体系搭建"},
    {"name": "刘璐莹", "role": "市场策略负责人 / 文档专员", "desc": "负责市场调研分析、商业计划书撰写、竞赛材料整合与品牌传播策略"},
]


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D7E2"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None):
    """创建带样式的表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = "微软雅黑"
        set_cell_shading(cell, "3DB8A0")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT
            run.font.name = "微软雅黑"
            if c_idx == 0:
                run.bold = True
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F4F7FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        if level == 1:
            run.font.color.rgb = ACCENT
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = DARK_TEXT
            run.font.size = Pt(14)
        elif level == 3:
            run.font.color.rgb = DARK_TEXT
            run.font.size = Pt(12)
    return h


def add_para(doc, text, bold=False, size=10.5, color=DARK_TEXT, align=None, space_after=6, first_line_indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return p


def add_bullet(doc, text, level=0, bold_prefix=""):
    """添加项目符号段落"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))

    # Clear default and re-add
    p.clear()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = "微软雅黑"
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = DARK_TEXT
        run_b.bold = True
        run_t = p.add_run(text)
        run_t.font.name = "微软雅黑"
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = DARK_TEXT
    else:
        run = p.add_run(text)
        run.font.name = "微软雅黑"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT
    return p


# ═══════════════════════════════════════════════════════════════
# MAIN DOCUMENT
# ═══════════════════════════════════════════════════════════════

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── 页眉页脚 ──
for section in doc.sections:
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("SparkOrbit 星轨学图 — 商业计划书（2026 iFLYTEK AI 开发者大赛 · 爱拼才慧莹）")
    hr.font.size = Pt(8)
    hr.font.color.rgb = GRAY_TEXT
    hr.font.name = "微软雅黑"

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("保密文件 · 吉林外国语大学")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY_TEXT
    fr.font.name = "微软雅黑"

# ═══════════════════════════════════════
# 封面
# ═══════════════════════════════════════

# 空行撑满
for _ in range(6):
    add_para(doc, "", size=12, space_after=0)

# 主标题
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(8)
r_t = p_title.add_run("SparkOrbit 星轨学图")
r_t.font.name = "微软雅黑"
r_t.font.size = Pt(32)
r_t.font.color.rgb = ACCENT
r_t.bold = True

# 副标题
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
r_s = p_sub.add_run("多智能体协同的个性化高等教育学习平台")
r_s.font.name = "微软雅黑"
r_s.font.size = Pt(16)
r_s.font.color.rgb = DARK_TEXT

# 英文副标题
p_en = doc.add_paragraph()
p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_en.paragraph_format.space_after = Pt(24)
r_e = p_en.add_run("Cognitive-Twin-Driven Adaptive Learning Ecosystem for Higher Education")
r_e.font.name = "Consolas"
r_e.font.size = Pt(10)
r_e.font.color.rgb = GRAY_TEXT
r_e.italic = True

# 分隔线
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_l = p_line.add_run("━" * 50)
r_l.font.color.rgb = ACCENT
r_l.font.size = Pt(8)

for _ in range(2):
    add_para(doc, "", size=10, space_after=0)

# 信息区
info_items = [
    ("项目名称", "SparkOrbit 星轨学图"),
    ("参赛赛道", "2026 iFLYTEK AI 开发者大赛 — Adaptive-LPDS（自适应个性化学习系统）"),
    ("团队名称", TEAM_NAME),
    ("所属院校", SCHOOL),
    ("提交日期", "2026 年 8 月"),
    ("演示站点", "https://wikj.online"),
]

for label, value in info_items:
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(4)
    r_label = p_info.add_run(f"{label}：")
    r_label.font.name = "微软雅黑"
    r_label.font.size = Pt(11)
    r_label.font.color.rgb = GRAY_TEXT
    r_label.bold = True
    r_val = p_info.add_run(value)
    r_val.font.name = "微软雅黑"
    r_val.font.size = Pt(11)
    r_val.font.color.rgb = DARK_TEXT

for _ in range(3):
    add_para(doc, "", size=10, space_after=0)

# 版本声明
p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_v = p_ver.add_run("本文件为保密商业计划书，仅供大赛评审及合作洽谈使用，未经授权不得外传。")
r_v.font.name = "微软雅黑"
r_v.font.size = Pt(8)
r_v.font.color.rgb = GRAY_TEXT
r_v.italic = True

doc.add_page_break()

# ═══════════════════════════════════════
# 目录占位
# ═══════════════════════════════════════
add_heading(doc, "目  录", level=1)
toc_items = [
    "一、执行摘要",
    "二、市场分析",
    "    2.1 宏观环境分析（PEST）",
    "    2.2 市场规模与趋势",
    "    2.3 目标用户画像",
    "    2.4 竞争格局分析",
    "    2.5 SWOT 分析",
    "三、产品与解决方案",
    "    3.1 产品定位",
    "    3.2 核心功能矩阵",
    "    3.3 技术创新亮点",
    "    3.4 防幻觉机制",
    "    3.5 产品路线图",
    "四、商业模式",
    "    4.1 商业画布",
    "    4.2 收入模型",
    "    4.3 定价策略",
    "    4.4 增长飞轮",
    "五、营销与拓展策略",
    "六、技术实现与团队",
    "七、运营与落地计划",
    "八、财务预测",
    "九、附录",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(item)
    r.font.name = "微软雅黑"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_TEXT if not item.startswith("    ") else GRAY_TEXT

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 一、执行摘要
# ═══════════════════════════════════════════════════════════
add_heading(doc, "一、执行摘要", level=1)

add_para(doc,
    "SparkOrbit 星轨学图是一个以「认知孪生」为核心、以「星系隐喻」为产品语言的沉浸式高等教育个性化学习平台。"
    "项目面向人工智能、计算机、电子信息等理工科专业的大班教学场景，通过对话式六维画像构建学生数字镜像，"
    "利用多智能体协同预演（Teacher → Mirror → Evaluator → PathPlanner）在真实作答前暴露认知误区、生成个性化补救路径，"
    "形成「画像—预演—路径—评估—教师复核」的完整学习闭环。",
    first_line_indent=0.74)

add_heading(doc, "1.1 核心价值主张", level=2)

add_para(doc,
    "区别于传统 LMS 平台的「资源搬运 + 进度记录」模式和通用 AI 学伴的「单聊问答」形态，SparkOrbit 提供三大差异化价值：",
    first_line_indent=0.74)

add_bullet(doc, "：建立包含专业背景、前置知识、认知风格、易错倾向、学习目标、时间弹性的六维画像，支持缺维追问与事件驱动增量刷新，将学生从「统一进度条」升级为「可计算的数字孪生」。", bold_prefix="认知孪生画像")
add_bullet(doc, "：基于 LangGraph StateGraph 编排四角色仿真流水线，在正式刷题前以学生数字孪生体试错，提前暴露误区并通过 SSE 流式展示推理过程，显著减少无效练习——这是目前国内同类产品中较罕见的工程化落地形态。", bold_prefix="多智能体预演")
add_bullet(doc, "：AI 产出（画像维度评分、改进建议、判题结果）均保留教师覆盖与工单裁决入口，避免算法黑箱替代教学决策，呼应教育场景对「人机协同」的根本要求。", bold_prefix="人机协同决策")

add_heading(doc, "1.2 市场机会", level=2)

add_para(doc,
    "据教育部《教育信息化 2.0 行动计划》与中国信通院《智慧教育发展报告》，中国智慧教育市场规模已突破 6000 亿元人民币，"
    "其中高等教育信息化占比约 18%，年复合增长率超过 12%。与此同时，高校大班教学的个性化困境日益突出——一名专业课教师平均面对 80-150 名学生，"
    "传统 LMS 仅能记录进度与分数，无法刻画个体认知差异。AI 大语言模型与多模态能力的成熟，使得「以学生认知状态为中心」的自适应学习成为可工程落地的新兴赛道。"
    "我们估算，以全国约 2000 所高校中理工科专业为切入目标市场，仅 B2B 私有化部署的 TAM（总可寻址市场）即超过 50 亿元人民币。",
    first_line_indent=0.74)

add_heading(doc, "1.3 商业模式与融资需求", level=2)

add_para(doc,
    "平台采用「B2B2C + B2C SaaS」双轨商业模式。B2B2C 面向高校提供私有化部署年费（按院系/学生规模阶梯定价），"
    "B2C 面向学生提供基础星系免费 + 高级多智能体预演按 Token 订阅的个人服务。"
    "增值服务包括教师星系锻造工坊、校本 RAG 定制、跨校星系共创市场等。"
    "项目当前已完成 MVP 核心闭环开发，并在公网部署可演示版本（https://wikj.online）。计划融资 200-300 万元种子轮，"
    "用于校本 RAG 知识库扩充、移动端降级方案开发、首批 3-5 所合作院校试点推广及团队扩张。预计 18 个月内达到盈亏平衡。",
    first_line_indent=0.74)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 二、市场分析
# ═══════════════════════════════════════════════════════════
add_heading(doc, "二、市场分析", level=1)

add_heading(doc, "2.1 宏观环境分析（PEST）", level=2)

add_heading(doc, "2.1.1 政策环境（Political）", level=3)
add_para(doc,
    "教育部《教育信息化 2.0 行动计划》（教技〔2018〕6 号）明确提出「构建网络化、数字化、智能化、个性化、终身化的教育体系」；"
    "《关于推进教育新型基础设施建设构建高质量教育支撑体系的指导意见》（教发〔2021〕2 号）将「智慧教学设施」列为重点方向，"
    "强调「利用人工智能技术实现因材施教」。2023 年以来，《生成式人工智能服务管理暂行办法》等法规为 AI+教育划定了"
    "合规边界——这与 SparkOrbit 始终强调的「人机协同终裁、隐私本地化处理」的设计理念高度契合。",
    first_line_indent=0.74)

add_heading(doc, "2.1.2 经济环境（Economic）", level=3)
add_para(doc,
    "中国高等教育经费持续增长，2024 年全国一般公共预算教育支出超过 4 万亿元。高校信息化预算中，"
    "智慧教学平台的采购优先级逐年提升，尤其在后疫情时代混合教学常态化背景下，各校对于能够同时满足"
    "「线上自主学习 + 教师学情治理 + 教务管理」一体化需求的平台需求旺盛。",
    first_line_indent=0.74)

add_heading(doc, "2.1.3 社会环境（Social）", level=3)
add_para(doc,
    "Z 世代大学生作为数字原住民，对学习工具的游戏化体验、即时反馈与社交互动有较高期待。"
    "与此同时，学业压力引发的心理健康问题日益受关注，学习系统若仅关注「成绩」而忽视情绪支"
    "持与专注管理，难以获得长期用户粘性。SparkOrbit 的星语树洞、AI 陪伴、自习督导、休闲激励等"
    "多维分区设计正是对这一社会趋势的产品回应。",
    first_line_indent=0.74)

add_heading(doc, "2.1.4 技术环境（Technological）", level=3)
add_para(doc,
    "大语言模型（DeepSeek、讯飞星火等）推理成本的持续下降与多模态能力（语音评测、图生图、"
    "文生视频）的成熟，使「以多智能体协同完成个性化教学闭环」成为在工程上可行且在成本上可"
    "承受的方向。同时，LangGraph 等 Agent 编排框架的崛起，降低了多角色协作的工程复杂度；"
    "ChromaDB 等轻量向量数据库的普及，使校本级 RAG 在小规模部署中不再依赖昂贵的基础设施。",
    first_line_indent=0.74)

add_heading(doc, "2.2 市场规模与趋势", level=2)

add_para(doc,
    "根据艾瑞咨询、弗若斯特沙利文等机构发布的教育科技行业报告，中国智慧教育市场近年呈现以下趋势：",
    first_line_indent=0.74)

add_styled_table(doc,
    ["指标", "2023 年", "2024 年", "2025E", "2026E", "2027E"],
    [
        ["中国智慧教育市场规模（亿元）", "5,200", "5,850", "6,500", "7,200", "8,000"],
        ["其中高等教育信息化占比", "17%", "18%", "19%", "20%", "20%"],
        ["AI+教育工具渗透率", "8%", "14%", "22%", "32%", "40%"],
        ["个性化学习平台增速 YoY", "——", "25%", "30%", "35%", "35%"],
    ],
    col_widths=[4.5, 2.2, 2.2, 2.2, 2.2, 2.2]
)

add_para(doc, "")
add_para(doc,
    "我们的总可寻址市场（TAM）推算如下：全国约 2,700 所普通高等学校，其中开设理工科专业的院校约 2,000 所。"
    "以平均每校 3 个理工科学院、每院年均信息化采购预算 10-15 万元计算，B2B 私有化部署的年 TAM 约为 60-90 亿元。"
    "B2C SaaS 方面，以理工科在校生约 1,200 万人、5% 付费渗透率、年均 ARPU 120 元计算，年 TAM 约为 7.2 亿元。"
    "合计 TAM 约为 67-97 亿元/年，市场空间充裕。",
    first_line_indent=0.74)

add_heading(doc, "2.3 目标用户画像", level=2)

add_styled_table(doc,
    ["角色", "典型用户", "核心痛点", "SparkOrbit 解决方案"],
    [
        ["学生", "理工科本科生 / 研究生\n（AI/计算机/电子信息专业）",
         "画像仅记对错、路径千人一面\n补救滞后、情绪缺支持",
         "六维认知画像 + 多智能体预演\n个性化路径 + 星语树洞心理支持"],
        ["教师", "高校专业课教师\n（大班教学场景）",
         "难筛风险学生、难复核 AI 建议\n讲义难结构化为练习题",
         "学情热力图 + 风险看板 + 改进复核\n星系锻造（PDF→知识宇宙）"],
        ["管理员", "教务处 / 信息中心\n运维人员",
         "Token 用量不可见\n缺统一治理面板",
         "用量监控 + 异常告警\n维护模式一键启停"],
    ],
    col_widths=[2.5, 3.5, 4.5, 4.5]
)

add_heading(doc, "2.4 竞争格局分析", level=2)

add_para(doc,
    "当前高等教育个性化学习市场处于早期阶段，主要竞争力量可分为三类：传统 LMS 平台、新兴 AI 学伴工具、以及大厂教育 AI 产品。"
    "SparkOrbit 的差异化优势在于将「认知画像—多智能体预演—教师协同决策」形成完整闭环，而非仅在单一环节提供 AI 辅助。",
    first_line_indent=0.74)

add_styled_table(doc,
    ["对比维度", "传统 LMS\n（超星/雨课堂/学堂在线）", "AI 学伴\n（ChatGPT Edu/Khanmigo）", "SparkOrbit 星轨学图"],
    [
        ["学习者建模", "进度 / 正确率", "会话级，弱持久化", "持久六维画像 + 教师复核"],
        ["知识组织", "章节列表 / 视频目录", "无结构", "星系—行星—星座 + 衰减状态机"],
        ["诊断方式", "练后统计", "单次问答", "多智能体作答前预演暴露误区"],
        ["教师角色", "发公告 / 批改", "基本缺位", "看板 / 锻造 / TimeWarp / 改进覆盖"],
        ["资源生成", "静态课件上传", "文本问答", "6 类 Agent 流式生成 + 质量自动评分"],
        ["多模态", "少", "文本为主", "语音评测 + 短视频 + 视觉督导 + 数字分身"],
        ["防幻觉", "不涉及", "无专门机制", "Agent 隔离 + RAG + 低置信工单 + 教师终裁"],
        ["部署形态", "云端 SaaS / 私有化", "云端 API", "SaaS + Docker 私有化 + 本机一键启动"],
    ],
    col_widths=[2.5, 3.5, 3.5, 5.5]
)

add_heading(doc, "2.5 SWOT 分析", level=2)

swot_data = [
    ["优势 Strengths", "劣势 Weaknesses"],
    [
        "• 完整学习闭环：画像—预演—路径—评估\n"
        "• 多模型协同矩阵：按任务最优路由\n"
        "• 三角色深度覆盖：学生/教师/管理员\n"
        "• 游戏化+沉浸式 UI：星轨世界观分区设计\n"
        "• 已获可运行 MVP，公网部署可演示\n"
        "• 16 份国标软件工程文档齐套",

        "• 尚无 LMS/教务系统对接（规划中）\n"
        "• 移动端体验待优化\n"
        "• 品牌认知度为零\n"
        "• 种子数据积累不足\n"
        "• 未获付费验证\n"
        "• 校本 RAG 知识库覆盖度有限"
    ],
    ["机会 Opportunities", "威胁 Threats"],
    [
        "• 国家教育数字化战略持续加码\n"
        "• AI+教育政策红利窗口期\n"
        "• 讯飞开放平台生态资源（本赛道）\n"
        "• 高校智慧校园建设刚性需求\n"
        "• LLM 推理成本持续下降\n"
        "• 跨校知识共享市场蓝海",

        "• LLM 幻觉风险（教学场景零容忍）\n"
        "• 大厂入局（字节/阿里/腾讯）\n"
        "• 数据隐私法规趋严\n"
        "• 高校采购决策周期长（6-18 月）\n"
        "• 多模型 API 成本波动\n"
        "• 开源 LLM 可能挤压商业价值"
    ],
]

# SWOT 表格用两个并排的 2x2
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(table)

for r_idx in range(2):
    for c_idx in range(2):
        cell = table.rows[r_idx].cells[c_idx]
        cell.text = ""
        lines = swot_data[r_idx * 2 + c_idx][1].split("\n")
        title = swot_data[r_idx * 2 + c_idx][0]
        p_t = cell.paragraphs[0]
        r_t = p_t.add_run(title)
        r_t.bold = True
        r_t.font.size = Pt(10)
        r_t.font.name = "微软雅黑"
        if r_idx == 0 and c_idx == 0:
            r_t.font.color.rgb = WHITE
            set_cell_shading(cell, "3DB8A0")
        elif r_idx == 0 and c_idx == 1:
            r_t.font.color.rgb = WHITE
            set_cell_shading(cell, "E8A838")
        elif r_idx == 1 and c_idx == 0:
            r_t.font.color.rgb = WHITE
            set_cell_shading(cell, "1A2332")
        else:
            r_t.font.color.rgb = WHITE
            set_cell_shading(cell, "6B7C93")
        for line in lines:
            if line.strip():
                p_l = cell.add_paragraph(line.strip())
                p_l.paragraph_format.space_after = Pt(1)
                for run in p_l.runs:
                    run.font.size = Pt(8)
                    run.font.name = "微软雅黑"
                    run.font.color.rgb = DARK_TEXT

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 三、产品与解决方案
# ═══════════════════════════════════════════════════════════
add_heading(doc, "三、产品与解决方案", level=1)

add_heading(doc, "3.1 产品定位", level=2)

add_para(doc,
    "SparkOrbit 星轨学图定位于「认知孪生驱动的个性化高等教育 AI 学习平台」，以人工智能、计算机、"
    "电子信息相关专业课程为首批落地学科（响应赛题要求）。产品核心隐喻为「星系」：将学科知识组织为"
    "星系—行星—星座图谱，行星点亮代表掌握、衰减代表遗忘、超新星代表永久固化。"
    "六维认知画像（专业背景 / 前置知识 / 认知风格 / 易错倾向 / 学习目标 / 时间弹性）构成学生的「数字镜像」，"
    "驱动后续所有智能决策——资源生成、路径规划、评估报告、教师干预建议。",
    first_line_indent=0.74)

add_para(doc,
    "产品以科幻星轨 UI 降低学习门槛，以游戏化机制（行星点亮/桌宠/成就/星座解锁）维持长期动机，"
    "以教师工作台保证人工专业判断不被算法取代。一句话概括：「不是又一个 AI 答疑机器人，而是一个让 AI 帮你提前暴露盲区、"
    "让老师看得见每个学生的个性化学习操作系统。」",
    first_line_indent=0.74)

add_heading(doc, "3.2 核心功能矩阵", level=2)

add_heading(doc, "3.2.1 学生端：星轨领航台与六大分区", level=3)

add_styled_table(doc,
    ["分区", "功能要点", "服务维度"],
    [
        ["学习区 Learn",
         "Three.js 星系探索、行星挑战答题、AI 伴学（Companion/Tutor 苏格拉底/费曼）、"
         "教案、SOS、镜像/多元宇宙预演 SimulationConsole、学习路径、资源工坊 ResourceStudio（6 类 Agent）、"
         "错题本、智能测验、知识图谱、番茄钟等 17 项学习工具",
         "认知核心闭环"],
        ["我的星域 Domain",
         "成长总览、六维画像 MirrorDashboard（雷达图+趋势）、掌握度全景、"
         "称号/里程碑/成就墙、桌宠图鉴",
         "自我认知与复盘"],
        ["星语树洞 TreeHole",
         "AI 陪伴聊天、心情日记、匿名动态与评论互动",
         "情绪与心理支持"],
        ["聊天区 Chat",
         "班级/话题/群聊/私聊（WebSocket）、星愿墙、资料站（文件上传与分享）",
         "社交与协作"],
        ["自习区 Study",
         "3D 选房、专注计时、排行、环境音；摄像头 + TensorFlow.js coco-ssd 本地推理分心/离开提醒",
         "专注与自律"],
        ["休闲区 Leisure",
         "星球记忆翻牌、陨石躲避、星座连线、签到、桌宠养成、积分商城",
         "激励与留存"],
    ],
    col_widths=[2.5, 9, 3.5]
)

add_heading(doc, "3.2.2 教师端：班级治理与内容锻造", level=3)

add_styled_table(doc,
    ["模块", "功能说明"],
    [
        ["学情看板", "班级概览指标卡、掌握度热力图、引力陷阱（薄弱知识点）、风险学生分级列表"],
        ["作业 / 考勤 / 巡查", "作业布置与批改、成绩册（按学生/知识点双视角）、考勤记录、自习室在线巡查"],
        ["画像改进复核", "AI 三档预评（优秀/合格/不合格）→ 教师可覆盖评分，回写画像维度分数并落库可追溯"],
        ["AI 教案", "按行星（知识点）生成结构化教案（目标/重难点/例题/练习建议）"],
        ["星系锻造", "上传课程讲义 PDF → LLM 结构化抽取知识点 → 生成 Galaxy/Planet 节点 + 挑战题 → 发布到星图"],
        ["时空扭曲沙盘", "选择学生 → 手动覆盖画像维度 → 启动对照预演 → 对比默认与假设画像下的推演差异"],
        ["消息 / 资料", "班级/指定学生广播通知；教学资料与课件上传分发"],
    ],
    col_widths=[3.5, 11.5]
)

add_heading(doc, "3.2.3 管理端：系统运维与合规控制台", level=3)

add_styled_table(doc,
    ["模块", "功能说明"],
    [
        ["系统概览", "注册用户数、活跃用户、在线自习、当日 Token 消耗、LLM 调用成功率等运维仪表盘"],
        ["用户管理", "按角色/班级筛选；新建、编辑、停用、重置密码；RBAC 权限约束"],
        ["内容管理", "全平台星系/行星/AI 生成资源审核、上下架、屏蔽；配合 Shield 风控双重治理"],
        ["Token 用量监控", "按模型/用户/时间维度统计消耗与成本；阈值告警；成本控制依据"],
        ["接口异常", "聚合 ApiUsageLog 错误记录，按错误码/接口/用户聚类展示"],
        ["维护模式", "一键开启后非管理员写操作返回 503；版本发布或故障应急期间冻结业务写入"],
    ],
    col_widths=[3.5, 11.5]
)

add_heading(doc, "3.3 技术创新亮点", level=2)

add_bullet(doc, "（认知孪生工程化）：画像不只是展示雷达图，而是驱动仿真人格、路径与资源编排的统一数据源。"
    "六维字段对应教学决策所需信息（非泛化兴趣标签），支持缺维追问与学习事件增量刷新，降低冷启动失败率。")
add_bullet(doc, "（可观察多智能体流水线）：基于 LangGraph StateGraph 编排 Teacher → Mirror → Evaluator → PathPlanner 四角色协同预演。"
    "SSE 事件按角色流式推送，师生可观察 Teacher 陷阱设计、Evaluator 错因归因、Planner 补救步骤等完整思维链。支持 Multiverse（多元宇宙对照）与 TimeWarp（教师沙盘推演）两种扩展模式。")
add_bullet(doc, "（遗忘显性化状态机）：行星掌握度按艾宾浩斯启发实现五态迁移（lit → fading → meteor → dim → supernova）。"
    "复习可重置衰减，多次巩固标记为永久固化。将神经科学规律转化为可见的「陨石危机」与「超新星成就」产品叙事。")
add_bullet(doc, "（按任务路由的多模型协同矩阵）：DeepSeek 负责核心推理与质量评分，讯飞星火负责中文长文本与多轮对话，"
    "讯飞 IAT/ISE 负责普通话/英语语音评测，cantonese.ai 负责粤语，通义千问负责数字分身图生图，火山方舟 Seedance 1.0 Pro Fast 负责教学短视频生成。各引擎按最优任务匹配，兼顾质量与成本。")
add_bullet(doc, "（人机协同评分）：AI 对改进提交产出预评（优秀/合格/不合格），教师可在改进复核面板覆盖评分，"
    "覆盖动作落库可追溯。低置信判题（confidence < 0.55）或引用不一致时，自动生成「待人审」工单，由教师接管终裁。")

add_heading(doc, "3.4 防幻觉机制", level=2)

add_para(doc,
    "在教育场景中，「教错」比「不教」的危害更大。SparkOrbit 落地的防幻觉工程约束包括：",
    first_line_indent=0.74)

add_styled_table(doc,
    ["机制", "实现方式", "状态"],
    [
        ["Agent 角色隔离", "Teacher 出题与 Evaluator 判题分 Prompt/分职责；出题携带 knowledge_point_id + expected_key_points + traps", "已实现"],
        ["Evaluator 独立引用", "判题侧独立产出 cited_knowledge_point_id 与置信度（禁止抄出题 ID 自证）", "已实现"],
        ["低置信/矛盾转教师", "confidence < 0.55 或引用不一致时写入教师可见「待人审」工单", "已实现"],
        ["RAG 校本溯源", "ChromaDB 灌入 seed 行星大纲；Tutor / 判题注入检索上下文并返回 sources", "已实现"],
        ["资源质量自动评分", "六类资源生成后 DeepSeek 四维打分（准确性/画像贴合/完整性/幻觉风险）；低分自动重试", "已实现"],
        ["内容安全 Shield", "敏感词过滤 + 可选 LLM 审核，拦截有害输出", "已实现"],
        ["PDF 页码级强制引用", "出题与判题绑定 PDF 页码/段落，为每道题建立不可篡改的来源证据链", "规划中"],
    ],
    col_widths=[3.5, 8.5, 2]
)

add_heading(doc, "3.5 产品路线图", level=2)

add_styled_table(doc,
    ["阶段", "时间", "关键目标"],
    [
        ["MVP 完善期", "第 1-6 个月",
         "扩充校本 RAG 知识库（覆盖 3-5 门核心课程）；遗忘阈值开放为教师可调参数；"
         "移动端降级方案（星图 2D 模式）；启动首批 3-5 所高校试点"],
        ["产品化与商业化", "第 6-12 个月",
         "对接 LMS/排课系统花名册与成绩回写；建立预演 vs 真实作答对照数据集，量化诊断准确率；"
         "产出首批付费客户（≥ 2 所高校）；完成软件著作权登记"],
        ["规模化与生态", "第 12-18 个月",
         "跨校知识宇宙共享平台上线（隐私合规前提）；深化多模态诊断（手写公式识别、实验操作分析）；"
         "探索联邦学习/可信执行环境等隐私计算本地化部署方案；扩展至 20+ 合作院校"],
    ],
    col_widths=[2.5, 2, 10.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 四、商业模式
# ═══════════════════════════════════════════════════════════
add_heading(doc, "四、商业模式", level=1)

add_heading(doc, "4.1 商业画布（Business Model Canvas）", level=2)

canvas_data = [
    ["重要伙伴 KP", "关键业务 KA", "价值主张 VP", "客户关系 CR", "客户细分 CS"],
    [
        "• 科大讯飞开放平台\n  （赛道生态合作）\n"
        "• 高校信息化代理商\n"
        "• 云服务商\n  （阿里云/腾讯云）\n"
        "• 开源社区\n  （Vue/FastAPI 生态）\n"
        "• 教育研究机构",

        "• 产品持续迭代\n"
        "• 校本 RAG 知识库建设\n"
        "• 高校试点推广与运营\n"
        "• Agent 编排逻辑优化\n"
        "• 内容合规审核",

        "• 认知孪生画像：\n  六维刻画个体差异\n"
        "• 多智能体预演：\n  作答前暴露误区\n"
        "• 教师人机协同：\n  AI 建议+人类裁决\n"
        "• 星系知识组织：\n  游戏化降低认知负荷\n"
        "• 沉浸式多分区：\n  认知/情绪/社交全覆盖",

        "• 专属客户成功经理\n  （B2B 高校客户）\n"
        "• 在线社群与答疑\n"
        "• 教师培训认证体系\n"
        "• 学生积分激励计划\n"
        "• 年度用户大会",

        "B2B2C：\n"
        "• 高校理工科院系\n"
        "• 教务处/信息中心\n"
        "• 专业课教师\n\n"
        "B2C：\n"
        "• 理工科在校生\n"
        "• 考研/考证自学者\n"
        "• 终身学习者",
    ],
    [
        "", "核心资源 KR", "", "渠道通路 CH", "",
    ],
    [
        "",
        "• 多智能体仿真引擎\n"
        "• 六维画像模型\n"
        "• 校本 RAG 知识库\n"
        "• 多模型协同矩阵\n"
        "• 软件工程文档体系",

        "",

        "• 讯飞开放平台分发\n"
        "• 教育信息化展会\n"
        "• 高校信息化代理商\n"
        "• 学术论文与白皮书\n"
        "• 开源社区运营",

        "",
    ],
    ["成本结构 C$", "", "收入来源 R$", "", ""],
    [
        "• LLM API 调用费用（多模型矩阵）\n"
        "• 云服务器/带宽/存储\n"
        "• 人力成本（研发+运营+市场）\n"
        "• 市场推广与差旅\n"
        "• 合规与法务",

        "",

        "B2B2C 智慧校园：\n"
        "• 私有化部署年费\n"
        "  （¥8-25 万/院系/年）\n\n"
        "B2C SaaS：\n"
        "• 基础版免费\n"
        "• 高级多智能体 ¥19.9/月\n"
        "• 增值语料 ¥9.9/月\n\n"
        "增值服务：\n"
        "• 星系锻造工坊 ¥2,000/场\n"
        "• 校本 RAG 定制 ¥5 万起",

        "",
        "",
    ],
]

table = doc.add_table(rows=len(canvas_data), cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(table)

for r_idx, row_data in enumerate(canvas_data):
    for c_idx, val in enumerate(row_data):
        if not val or val == "":
            continue
        cell = table.rows[r_idx].cells[c_idx]
        if r_idx > 0 and c_idx > 0 and c_idx < 4 and canvas_data[r_idx - 1][c_idx] == canvas_data[r_idx][c_idx]:
            continue
        # Merge if same as above
        cell.text = ""
        lines = val.split("\n")
        first = True
        for line in lines:
            if not line.strip():
                continue
            if first:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line.strip())
                r.bold = True
                r.font.size = Pt(8)
                r.font.name = "微软雅黑"
                r.font.color.rgb = DARK_TEXT
                first = False
            else:
                p = cell.add_paragraph()
                r = p.add_run(line.strip())
                r.font.size = Pt(7.5)
                r.font.name = "微软雅黑"
                r.font.color.rgb = DARK_TEXT

# Set column widths approximately
for row in table.rows:
    for i, w in enumerate([3.5, 3.2, 4.2, 3.2, 3.5]):
        if i < len(row.cells):
            row.cells[i].width = Cm(w)

# Header shading
for c_idx in range(5):
    if table.rows[0].cells[c_idx].text.strip():
        set_cell_shading(table.rows[0].cells[c_idx], "3DB8A0")
        for p in table.rows[0].cells[c_idx].paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE

add_heading(doc, "4.2 收入模型", level=2)

add_para(doc,
    "平台采用双轨收入模型，兼顾 B 端高校采购的稳定现金流与 C 端个人订阅的增量空间：",
    first_line_indent=0.74)

add_heading(doc, "B2B2C 智慧校园私有化部署", level=3)
add_bullet(doc, "面向高校院系级采购，提供私有化部署（Docker Compose）或 VPC 隔离 SaaS 两种交付方式。")
add_bullet(doc, "年费按院系规模阶梯定价：小型（<500 学生）¥8 万/年，中型（500-2000 学生）¥15 万/年，大型（>2000 学生）¥25 万/年。")
add_bullet(doc, "含基础功能 + 星系锻造 + 学情看板；高级功能（TimeWarp、校本 RAG 定制）单独计费。")

add_heading(doc, "B2C SaaS 个人订阅", level=3)
add_bullet(doc, "基础星图 + 行星挑战 + 画像查看免费开放，作为获客漏斗。")
add_bullet(doc, "高级多智能体预演（Simulation）按 ¥19.9/月或 ¥199/年订阅（预演消耗大量 Token，需收费用以覆盖 API 成本）。")
add_bullet(doc, "增值语料库（高级发音测评、数字分身生成等）¥9.9/月。")

add_heading(doc, "增值服务", level=3)
add_bullet(doc, "教师星系锻造工坊：¥2,000/场（线上培训 + 1 门课程锻造辅导）。")
add_bullet(doc, "校本 RAG 定制：¥5 万起（根据本校教材/讲义定制专属向量知识库）。")
add_bullet(doc, "跨校星系共创市场：交易抽佣 15%（长期）。")

add_heading(doc, "4.3 定价策略", level=2)

add_styled_table(doc,
    ["版本", "适用对象", "年费", "核心功能", "亮点"],
    [
        ["基础版\n（自助）", "个人学习者", "免费", "星图探索 + 行星挑战 + 画像查看", "获客入口，零门槛体验"],
        ["高级版\n（订阅）", "重度学习者", "¥199/年", "无限多智能体预演 + 高级评测 + 分身", "模拟推算盲区，省刷题"],
        ["标准版\n（校本）", "中小院系\n<500 学生", "¥8 万/年", "私有化部署 + 班级治理 + 星系锻造", "讲义→知识宇宙，一次投入"],
        ["旗舰版\n（全校）", "大型院校\n>2000 学生", "¥25 万/年", "全功能 + RAG 定制 + TimeWarp + API 对接", "人机协同全覆盖"],
    ],
    col_widths=[2.5, 2.5, 2, 5, 3]
)

add_heading(doc, "4.4 增长飞轮", level=2)

add_para(doc,
    "SparkOrbit 的增长模型遵循数据网络效应的正向循环：",
    first_line_indent=0.74)

add_bullet(doc, "更多学生使用 → 画像数据积累 → AI 预演与路径更精准 → 学习效果提升 → 学生口碑传播 → 更多学生使用")
add_bullet(doc, "更多教师使用 → RAG 校本知识库扩充 → 出题/判题质量提升 → 教师信任增加 → 学校采购决策 → 更多教师使用")
add_bullet(doc, "更多学校部署 → 跨校星系共创 → 优质教学内容复用 → 新校部署成本降低 → 更多学校部署")
add_para(doc,
    "两个关键启动杠杆：(1) 通过讯飞大赛生态获取早期曝光与种子用户；(2) B2C 免费版扩大基数，B2B 私有化部署贡献主要收入，"
    "两者互为增长引擎。",
    first_line_indent=0.74)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 五、营销与拓展策略
# ═══════════════════════════════════════════════════════════
add_heading(doc, "五、营销与拓展策略", level=1)

add_heading(doc, "5.1 进入市场策略", level=2)

add_para(doc,
    "我们采用「试点验证 → 区域辐射 → 全国推广」三步走策略：",
    first_line_indent=0.74)

add_bullet(doc, "（0-6 个月）：依托大赛资源接触首批 3-5 所合作院校（优先吉林外国语大学本校及省内兄弟院校），"
    "以「AI/数据结构」课程为切入点完成教学闭环验证，收集师生使用反馈，产出教学效果对比数据与案例白皮书。", bold_prefix="种子试点期")
add_bullet(doc, "（6-12 个月）：在东北地区高校形成标杆案例后，通过教育信息化展会、高校 CIO 论坛等渠道向华北/华东区域辐射。"
    "与 1-2 家高校信息化代理商建立合作，加速覆盖。目标：签约 5-10 所付费院校。", bold_prefix="区域拓展期")
add_bullet(doc, "（12-18 个月）：跨校星系共创市场上线，以优质教学内容的网络效应推动自发增长。"
    "探索与讯飞教育业务线的深度合作，接入讯飞智慧教育解决方案体系，获得更大渠道推力。", bold_prefix="规模增长期")

add_heading(doc, "5.2 渠道策略", level=2)

add_styled_table(doc,
    ["渠道类型", "具体方式", "优先级"],
    [
        ["赛事生态", "借助 2026 iFLYTEK AI 开发者大赛获得专家评审反馈与讯飞平台曝光", "高"],
        ["学术影响力", "撰写 AI+教育技术白皮书；投稿 CCF/教育技术期刊；参与信息化教学研讨会分享案例", "中"],
        ["代理商网络", "与 2-3 家高校信息化集成商签订代理协议，利用其现有客户关系加速获客", "高"],
        ["开源社区", "将六维画像模型、LangGraph 编排模式等非业务核心模块开源，吸引开发者社区贡献", "中"],
        ["教师培训", "举办「AI 赋能教学」线上工作坊，以星系锻造为钩子吸引教师免费试用", "中"],
    ],
    col_widths=[3, 9.5, 1.5]
)

add_heading(doc, "5.3 品牌建设", level=2)

add_bullet(doc, "以「认知孪生」为核心概念打造差异化品牌认知——当教育科技行业谈及「自适应学习 2.0」时，首先想到 SparkOrbit。")
add_bullet(doc, "发布季度《高校 AI 学习行为报告》，基于脱敏数据输出行业洞察，建立思想领导力。")
add_bullet(doc, "在 B 站 / 知乎等平台运营「星轨学图」技术科普账号，以「让 AI 帮你发现你不知道自己不知道的」为内容主线。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 六、技术实现与团队
# ═══════════════════════════════════════════════════════════
add_heading(doc, "六、技术实现与团队", level=1)

add_heading(doc, "6.1 技术架构总览", level=2)

add_para(doc,
    "系统采用前后端分离的微服务化单体架构（Modular Monolith），兼顾开发效率与后续拆分的灵活性。",
    first_line_indent=0.74)

add_para(doc,
    "前端：Vue 3 + Vite + TypeScript + Pinia。学生端采用单页 Zone 状态机架构（六分区按需挂载，降低内存占用）；"
    "三角色通过路由守卫分流（meta.role），越权访问在前端与后端双层拦截。视觉采用 Tailwind CSS + Three.js（3D 星图）+ "
    "GSAP（微交互动效）+ ECharts（数据可视化）+ TensorFlow.js + coco-ssd（自习本地推理）。"
    "内容渲染基于 markdown-it + highlight.js，支持公式/表格/代码高亮。",
    first_line_indent=0.74)

add_para(doc,
    "后端：FastAPI 单应用承载全部业务（REST + WebSocket + SSE），统一前缀 /api。"
    "数据层基于 SQLAlchemy Async + MySQL（可切换 SQLite/PostgreSQL），认证采用轻量 Bearer Token。"
    "智能层以多模型协同矩阵为核心，按任务类型路由至最优引擎。仿真编排基于 LangGraph StateGraph。",
    first_line_indent=0.74)

add_para(doc,
    "部署：支持 Docker Compose 一键部署与 Windows 本机 start.bat 一键启动（Python + SQLite），降低评委验收与潜在客户试用的门槛。"
    "公网演示站 https://wikj.online 已上线可访问。",
    first_line_indent=0.74)

add_heading(doc, "6.2 核心技术栈总览", level=2)

add_styled_table(doc,
    ["层级", "技术选型", "用途"],
    [
        ["前端框架", "Vue 3 + Vite + TypeScript", "三角色单仓前端"],
        ["状态管理", "Pinia", "全局状态与 Zone 局部状态分离"],
        ["3D / 可视化", "Three.js / ECharts", "星图渲染与数据图表"],
        ["动画", "GSAP", "过渡与微交互动效"],
        ["前端 ML", "TensorFlow.js + coco-ssd", "自习督导本地目标检测"],
        ["后端框架", "FastAPI + Uvicorn", "REST / WebSocket / SSE 统一服务"],
        ["ORM / DB", "SQLAlchemy Async + MySQL", "关系数据持久化"],
        ["向量检索", "ChromaDB", "校本 RAG 知识库"],
        ["Agent 编排", "LangGraph (StateGraph)", "仿真四角色流水线"],
        ["核心推理", "DeepSeek", "评估归因 / 路径规划 / 画像抽取 / 质量评分"],
        ["中文生成", "讯飞星火 X2 / 4.0 Turbo", "文档/题库生成 / Companion 伴学"],
        ["语音评测", "讯飞 IAT / ISE", "普通话 & 英语 ASR + 口语评测"],
        ["粤语", "cantonese.ai", "粤语 STT + 发音评分"],
        ["图像生成", "通义千问 Qwen (DashScope)", "数字分身图生图"],
        ["视频生成", "火山方舟 Seedance 1.0 Pro Fast", "教学短视频生成（失败降级 GSAP）"],
    ],
    col_widths=[3, 5.5, 6.5]
)

add_heading(doc, "6.3 已实现能力与工程成熟度", level=2)

add_para(doc,
    "当前 MVP 已完成赛题全部核心功能要求，以下为关键能力实现状态：",
    first_line_indent=0.74)

add_styled_table(doc,
    ["能力项", "状态", "说明"],
    [
        ["对话式六维画像 + 事件刷新", "已实现", "profiling / profile_refresh；挑战/费曼后增量更新"],
        ["6 类 Resource Agent", "已实现", "Doc / Mind / Quiz / Read / Media / Code"],
        ["Seedance 教学短视频", "已实现", "已开通接入点；异步生成并落盘；失败降级 GSAP/缓存"],
        ["LangGraph 仿真编排", "已实现", "Teacher→Mirror→Evaluator→PathPlanner + SSE 流式"],
        ["苏格拉底 / 费曼 Tutor", "已实现", "苏格拉底默认开启 + 费曼模式可选"],
        ["学习路径规划 + 评估回灌", "已实现", "画像驱动步骤生成 + 评估报告一键重排路径"],
        ["低置信/矛盾工单 + 教师复核", "已实现", "Evaluator 独立引用 + confidence < 0.55 转教师"],
        ["校本 RAG (ChromaDB)", "已实现", "种子行星大纲灌入；无命中回退伪源保证演示不空"],
        ["资源质量自动评分 (A/P/C/H)", "已实现", "DeepSeek 四维打分；低分自动重试"],
        ["Shield 内容安全", "已实现", "敏感词过滤 + 可选 LLM 审核"],
        ["自习督导本地推理", "已实现", "TF.js + coco-ssd；视频流不上传服务器"],
        ["三角色完整覆盖", "已实现", "学生/教师/管理员 + RBAC 双层守卫"],
        ["16 份国标软件工程文档", "已完成", "A-G 全生命周期文档齐套"],
    ],
    col_widths=[5.5, 1.5, 8]
)

add_heading(doc, "6.4 团队介绍", level=2)

add_para(doc,
    f"SparkOrbit 由 {SCHOOL} {TEAM_NAME} 团队开发。团队虽精干但分工明确，覆盖了从系统架构到产品设计、从技术实现到商业策略的全链路能力：",
    first_line_indent=0.74)

for m in MEMBERS:
    add_bullet(doc, "", bold_prefix=f"{m['name']} — {m['role']}：")
    add_bullet(doc, m['desc'], level=1)

add_heading(doc, "6.5 知识产权策略", level=2)

add_bullet(doc, "项目源代码已于竞赛提交时完整提交，核心创新点（多智能体仿真编排方法、基于六维画像的个性化学习路径生成算法）拟申请软件著作权及发明专利。")
add_bullet(doc, "前端采用开源技术栈（Vue3/Vite/Three.js/GSAP/TensorFlow.js），遵守各开源协议；后端核心逻辑自主开发，无第三方代码版权风险。")
add_bullet(doc, "品牌名称「SparkOrbit 星轨学图」及星轨 Logo 拟注册商标。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 七、运营与落地计划
# ═══════════════════════════════════════════════════════════
add_heading(doc, "七、运营与落地计划", level=1)

add_heading(doc, "7.1 实施路线图（18 个月）", level=2)

add_styled_table(doc,
    ["时间", "产品与技术", "商业与运营", "关键里程碑"],
    [
        ["第 1-3 个月\nMVP 完善",
         "• 扩充校本 RAG（覆盖 3 门核心课程）\n"
         "• 遗忘阈值班级可配\n"
         "• 移动端星图 2D 降级方案\n"
         "• LLM 调用成本持续优化",
         "• 大赛材料终稿与路演准备\n"
         "• 接触 5 所目标院校\n"
         "• 发布技术白皮书 v1.0\n"
         "• 开源六维画像模型",
         "大赛提交完成\n获得评委反馈"],
        ["第 4-6 个月\n试点启动",
         "• PDF 页码级强制引用（防幻觉增强）\n"
         "• 首批试点院校部署\n"
         "• 教师培训工坊 x2 场",
         "• 签约 2 所试点院校\n"
         "• 收集 >500 学生行为数据\n"
         "• 产出教学效果对比报告",
         "首批试点数据回收\nNPS ≥ 40"],
        ["第 7-9 个月\n商业化准备",
         "• LMS/教务系统对接（API 适配）\n"
         "• 预演 vs 真实作答校准\n"
         "• 付费墙与订阅系统开发",
         "• 参加 2 场教育信息化展会\n"
         "• 签约 1 家代理商\n"
         "• 启动 B2C 免费版公测",
         "B2B 首单签约\nB2C 注册 >2000"],
        ["第 10-12 个月\n付费验证",
         "• Agent 可解释报告上线\n"
         "• 性能压测与稳定性优化\n"
         "• 安全审计与渗透测试",
         "• 目标：5 所付费院校\n"
         "• B2C 付费转化率 ≥ 3%\n"
         "• 软件著作权申请提交",
         "月经常性收入 MRR ≥ ¥10 万\n软著受理"],
        ["第 13-18 个月\n规模化",
         "• 跨校星系共创市场\n"
         "• 多模态诊断深化\n"
         "• 隐私计算本地化部署试点",
         "• 合作院校扩展至 20+\n"
         "• 区域代理商 ≥ 3 家\n"
         "• 种子轮融资完成",
         "月经常性收入 MRR ≥ ¥50 万\n盈亏平衡达成"],
    ],
    col_widths=[2.5, 5.5, 5, 3]
)

add_heading(doc, "7.2 关键绩效指标（KPI）", level=2)

add_styled_table(doc,
    ["指标类别", "KPI", "6 个月目标", "12 个月目标", "18 个月目标"],
    [
        ["用户规模", "活跃学校数", "2 所", "8 所", "25 所"],
        ["用户规模", "日活学生数 (DAU)", "300", "2,000", "8,000"],
        ["用户规模", "活跃教师数", "10", "60", "200"],
        ["产品使用", "行星挑战完成量/日", "150", "1,500", "8,000"],
        ["产品使用", "多智能体预演次数/日", "30", "400", "2,500"],
        ["商业表现", "B2B 年化合同额 (ARR)", "——", "¥60 万", "¥400 万"],
        ["商业表现", "B2C 付费用户数", "——", "300", "2,500"],
        ["商业表现", "月经常性收入 (MRR)", "——", "¥10 万", "¥50 万"],
    ],
    col_widths=[2.5, 4.5, 2.5, 2.5, 3]
)

add_heading(doc, "7.3 风险与应对策略", level=2)

add_styled_table(doc,
    ["风险类别", "具体风险", "影响程度", "发生概率", "应对措施"],
    [
        ["技术风险", "LLM 幻觉导致错误教学", "高", "中",
         "Agent 隔离 + RAG 溯源 + 低置信转教师工单 + 资源质量自动评分（已落地四重防线）"],
        ["技术风险", "API 成本随用量激增", "高", "高",
         "大小模型按任务路由（核心推理走 DeepSeek，闲聊走轻量接口）；计划引入语义缓存与限流削峰"],
        ["技术风险", "第三方服务宕机导致降级", "中", "中",
         "多模型互为备份；Seedance 失败降级 GSAP/缓存片；核心答题链路不依赖单一模型"],
        ["市场风险", "大厂推出竞品", "中", "高",
         "深耕高校垂直场景与教师工作流，以场景深度而非模型能力竞争；锁定早期合作院校建立转换成本"],
        ["市场风险", "高校采购周期长（6-18 月）", "高", "高",
         "以 B2C 免费版先行获客，积累数据与口碑；提供试用期与效果评估报告降低采购决策门槛"],
        ["合规风险", "数据隐私法规趋严", "高", "中",
         "自习督导全程本地推理（视频不出浏览器）；数据最小化采集；规划引入监护人同意流程与独立审计库"],
        ["运营风险", "种子用户激活与留存不足", "中", "中",
         "以教师星系锻造为钩子驱动首次使用；游戏化机制（行星点亮/桌宠/成就）+ 班级社交提升粘性"],
    ],
    col_widths=[2, 3.5, 1.5, 1.5, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 八、财务预测
# ═══════════════════════════════════════════════════════════
add_heading(doc, "八、财务预测", level=1)

add_para(doc,
    "以下财务预测基于保守假设，未考虑跨校星系共创市场交易抽佣等远期收入。所有数据为预估值，实际结果可能因市场环境、执行效率等因素产生偏差。",
    first_line_indent=0.74, color=GRAY_TEXT, size=9)

add_heading(doc, "8.1 收入预测（3 年）", level=2)

add_styled_table(doc,
    ["收入来源", "第 1 年", "第 2 年", "第 3 年", "备注"],
    [
        ["B2B 私有化部署", "¥60 万", "¥280 万", "¥700 万",
         "按 8 万/校院起，第 2 年起标杆案例带动价格上浮"],
        ["B2C 个人订阅", "¥7 万", "¥50 万", "¥180 万",
         "B2C 前 6 个月免费获客，随后 ¥199/年定价"],
        ["增值服务", "¥3 万", "¥30 万", "¥80 万",
         "教师培训工坊 + 校本 RAG 定制"],
        ["合计收入", "¥70 万", "¥360 万", "¥960 万", ""],
    ],
    col_widths=[3.5, 2.5, 2.5, 2.5, 4]
)

add_heading(doc, "8.2 成本结构", level=2)

add_styled_table(doc,
    ["成本项目", "第 1 年", "第 2 年", "第 3 年", "说明"],
    [
        ["人力成本", "¥36 万", "¥96 万", "¥180 万",
         "3 人→5 人→8 人；含社保"],
        ["LLM API 调用", "¥8 万", "¥35 万", "¥80 万",
         "随用户量增长；路由优化后边际递减"],
        ["云服务器/带宽", "¥4 万", "¥15 万", "¥35 万",
         "含私有化部署支持"],
        ["市场推广", "¥5 万", "¥20 万", "¥50 万",
         "展会/代理商佣金/线上投放"],
        ["其他（合规/法务/差旅）", "¥3 万", "¥8 万", "¥20 万", ""],
        ["合计成本", "¥56 万", "¥174 万", "¥365 万", ""],
    ],
    col_widths=[3.5, 2.5, 2.5, 2.5, 4]
)

add_heading(doc, "8.3 盈亏平衡分析", level=2)

add_styled_table(doc,
    ["指标", "第 1 年", "第 2 年", "第 3 年"],
    [
        ["收入", "¥70 万", "¥360 万", "¥960 万"],
        ["成本", "¥56 万", "¥174 万", "¥365 万"],
        ["毛利", "¥14 万", "¥186 万", "¥595 万"],
        ["毛利率", "20%", "52%", "62%"],
        ["累计利润", "¥14 万", "¥200 万", "¥795 万"],
    ],
    col_widths=[4, 3.5, 3.5, 3.5]
)

add_para(doc, "")
add_para(doc,
    "预计第 10-12 个月实现单月盈亏平衡（MRR ≈ ¥10 万时覆盖月均运营成本 ¥7-8 万）。"
    "第 2 年起毛利率稳步提升，主要受益于：(1) LLM 调用成本随路由优化与批量折扣下降；"
    "(2) B2B 标准产品化率提高，私有化部署的人均支持成本摊薄；(3) B2C 用户基数扩大后的边际成本递减效应。",
    first_line_indent=0.74)

add_heading(doc, "8.4 融资计划", level=2)

add_para(doc,
    "项目计划在第一年进行种子轮融资，用于加速产品迭代与市场验证：",
    first_line_indent=0.74)

add_styled_table(doc,
    ["轮次", "目标金额", "时间节点", "资金用途"],
    [
        ["种子轮", "¥200-300 万", "第 1 年 Q2-Q3\n（大赛结束后 3-6 个月）",
         "校本 RAG 知识库扩充（30%）、移动端降级方案开发（20%）、"
         "首批 3-5 所合作院校试点推广（25%）、团队扩张至 5 人（25%）"],
        ["天使轮", "¥500-800 万", "第 2 年 Q3-Q4\n（付费验证后）",
         "LMS 对接与 API 生态建设（30%）、跨校星系共创市场开发（25%）、"
         "市场推广与代理商体系建设（25%）、隐私计算本地化部署研发（20%）"],
    ],
    col_widths=[2, 2.5, 3, 7.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 九、附录
# ═══════════════════════════════════════════════════════════
add_heading(doc, "九、附录", level=1)

add_heading(doc, "9.1 软件工程文档清单", level=2)

add_para(doc,
    "本项目已完成 16 份国标软件工程文档（A-G 生命周期），彰显工程规范性与交付成熟度：",
    first_line_indent=0.74)

add_styled_table(doc,
    ["文档类别", "文档名称", "状态"],
    [
        ["A 可行性", "A1 可行性研究报告、A2 项目开发计划", "已完成"],
        ["B 需求", "B1 软件需求说明书、B2 数据要求说明书", "已完成"],
        ["C 设计", "C1 概要设计说明书、C2 详细设计说明书、C3 数据库设计说明书", "已完成"],
        ["D 实现", "D1 模块开发卷宗、D2 用户手册、D3 操作手册", "已完成"],
        ["E 测试", "E1 测试计划、E2 测试分析报告", "已完成"],
        ["F 运维", "F1 开发进度月报（含 2 期）、F2 项目开发总结报告", "已完成"],
        ["G 质量", "G1 软件质量保证计划、G2 软件配置管理计划", "已完成"],
        ["H 竞赛", "作品设计实现方案、部署说明书、演示视频脚本、PPT 方案", "已完成"],
    ],
    col_widths=[2, 8, 2]
)

add_heading(doc, "9.2 演示与验证信息", level=2)

add_styled_table(doc,
    ["项目", "内容"],
    [
        ["公网演示站", "https://wikj.online"],
        ["演示账号", "学生 student001 / 教师 teacher001 / 管理员 admin001（密码 123456）"],
        ["本机启动", "解压源码 → 双击 start.bat → 浏览器打开 http://127.0.0.1:8000"],
        ["Docker 部署", "docker-compose up -d（需 MySQL）"],
        ["OpenAPI 文档", "https://wikj.online/docs"],
        ["评分证据包", "docs/evidence/（资源/路径/辅导/评估案例 + 截图）"],
    ],
    col_widths=[3, 12]
)

add_heading(doc, "9.3 竞赛评分自评对照", level=2)

add_styled_table(doc,
    ["评分项", "占比", "自评得分", "核心依据"],
    [
        ["创新性", "20%", "17/20", "认知孪生 + 多智能体可观察预演 + 人机协同"],
        ["实用性", "15%", "13/15", "AI/计算机课程切入；星系锻造校本落地"],
        ["核心功能完成度", "25%", "23/25", "六维画像 + ≥6 种资源 + 路径 + 苏格拉底/费曼 + 评估"],
        ["技术实现质量", "15%", "13/15", "FastAPI + SSE + LangGraph + 多模型矩阵"],
        ["功能创新", "5%", "5/5", "Mirror 镜像预演 + TimeWarp + 星系锻造"],
        ["配套文档", "10%", "9/10", "16 份国标文档 + 设计实现方案 + 证据包"],
        ["演示视频/PPT", "10%", "8/10", "实操脚本完整；成片与终稿待固化"],
    ],
    col_widths=[4, 1.5, 2, 7.5]
)

add_para(doc, "")
add_para(doc, "自评合计：约 88/100 分（基础项），附加分约 8/10 分，总计约 96/110 分，折合百分制约 87 分。", bold=True)
add_para(doc, "视频/PPT 终稿补齐后可冲击 90+ 分。", color=GRAY_TEXT, size=9)

add_heading(doc, "9.4 联系方式", level=2)

add_para(doc, f"团队名称：{TEAM_NAME}", size=10)
add_para(doc, f"所属院校：{SCHOOL}", size=10)
add_para(doc, "演示站点：https://wikj.online", size=10)
add_para(doc, "赛道名称：2026 iFLYTEK AI 开发者大赛 — Adaptive-LPDS（自适应个性化学习系统）", size=10)
add_para(doc, "提交日期：2026 年 8 月", size=10)

# ── 保存 ──
doc.save(OUT_PATH)
print(f"商业计划书已生成: {OUT_PATH}")
print(f"文件大小: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
