# -*- coding: utf-8 -*-
"""验证最终文档"""
import docx

doc = docx.Document(r"c:\Users\咸\Desktop\软件杯\17016457介绍\项目设计文档_更新版.docx")

# Check preserved content
preserved = {
    "LangGraph": 0,
    "高等教育": 0,
    "人工智能": 0,
}
# Check updated content
updated = {
    "自研多智能体编排": 0,
    "中小学教育": 0,
    "火山方舟豆包": 0,
    "Tailwind CSS 4": 0,
    "数字人导师": 0,
}
# Check new content
new_content = {
    "四闸冲刺体系": 0,
    "练习中": 0,
    "exploring": 0,
    "引力黑洞": 0,
    "bootstrap_from_assessment": 0,
    "班级门控策略": 0,
    "GatePolicy": 0,
    "分镜讲稿": 0,
    "GSAP 动画逐幕": 0,
    "演武舱": 0,
    "Pyodide": 0,
    "口语训练舱": 0,
    "六类独立训练舱": 0,
    "关卡策略面板": 0,
    "分层教学": 0,
    "ZoneDock 21": 0,
    "数字人讲解": 0,
}

for para in doc.paragraphs:
    t = para.text
    for key in preserved:
        if key in t:
            preserved[key] += 1
    for key in updated:
        if key in t:
            updated[key] += 1
    for key in new_content:
        if key in t:
            new_content[key] += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            ct = cell.text
            for key in preserved:
                if key in ct:
                    preserved[key] += 1
            for key in updated:
                if key in ct:
                    updated[key] += 1
            for key in new_content:
                if key in ct:
                    new_content[key] += 1

print("=== SHOULD BE PRESERVED ===")
for k, v in preserved.items():
    status = "OK" if v > 0 else "MISSING!"
    print(f"  {k}: {v} ({status})")

print("\n=== SHOULD BE UPDATED ===")
for k, v in updated.items():
    if k == "自研多智能体编排":
        status = "OK (should be 0)" if v == 0 else "ERROR! should be 0"
    elif k == "中小学教育":
        status = "OK (should be 0)" if v == 0 else "ERROR! should be 0"
    else:
        status = "OK" if v > 0 else "MISSING!"
    print(f"  {k}: {v} ({status})")

print("\n=== NEW CONTENT SHOULD EXIST ===")
for k, v in new_content.items():
    if k == "ZoneDock 21":
        status = "OK" if v > 0 else "MISSING!"
    elif "GSAP" in k:
        status = "OK" if v > 0 else "MISSING!"
    elif "Pyodide" in k:
        status = "OK" if v > 0 else "MISSING!"
    else:
        status = "OK" if v > 0 else "MISSING!"
    print(f"  {k}: {v} ({status})")

# Spot check key paragraphs
print("\n=== SPOT CHECK ===")
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if "2.3.1.1" in t:
        print(f"P{i}: 2.3.1.1 Heading found!")
    if "四闸冲刺" in t and len(t) < 60:
        print(f"P{i}: {t}")
    if "引力黑洞" in t and len(t) < 80:
        print(f"P{i}: {t}")
    if "班级门控" in t and len(t) < 80:
        print(f"P{i}: {t}")
    if "分镜讲稿" in t and len(t) < 80:
        print(f"P{i}: {t}")
    if "关卡策略面板" in t:
        print(f"P{i}: {t}")

# Check toolbox table size
if len(doc.tables) > 8:
    t8 = doc.tables[8]
    print(f"\nTable 8 rows: {len(t8.rows)}")
    for ri in [1, 2, 3, -1]:
        cells = [c.text[:30] for c in t8.rows[ri].cells]
        idx = ri if ri >= 0 else len(t8.rows) + ri
        print(f"  Row {idx}: {cells}")

# Check Table 5
if len(doc.tables) > 5:
    t5 = doc.tables[5]
    print(f"\nTable 5 rows: {len(t5.rows)}")
    print(f"  Last row: {[c.text[:40] for c in t5.rows[-1].cells]}")
