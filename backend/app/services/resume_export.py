"""高级简历导出：HTML（打印/下载）/ Word / Markdown。"""
from __future__ import annotations

import base64
import html
import re
from io import BytesIO
from typing import Any

from app.data.career_templates import get_resume_template, resolve_template_id

_PHOTO_MAX = 1_800_000


def _lines(value: Any) -> list[str]:
    if isinstance(value, list):
        out = []
        for item in value:
            if isinstance(item, dict):
                text = " / ".join(str(item.get(k) or "") for k in ("name", "org", "role", "highlight") if item.get(k))
            else:
                text = str(item)
            if text.strip():
                out.append(text.strip())
        return out[:12]
    text = str(value or "")
    return [ln.strip() for ln in text.replace("\r", "").split("\n") if ln.strip()][:12]


def normalize_fields(fields: dict[str, Any] | None) -> dict[str, Any]:
    raw = fields or {}
    return {
        "name": str(raw.get("name") or "姓名").strip() or "姓名",
        "intent": str(raw.get("intent") or "").strip(),
        "city": str(raw.get("city") or "").strip(),
        "contact": str(raw.get("contact") or "").strip(),
        "email": str(raw.get("email") or "").strip(),
        "github": str(raw.get("github") or "").strip(),
        "education": _lines(raw.get("education")),
        "experience": _lines(raw.get("experience")),
        "projects": _lines(raw.get("projects")),
        "skills": _lines(raw.get("skills")),
        "certificates": _lines(raw.get("certificates")),
        "papers": _lines(raw.get("papers")),
        "highlights": _lines(raw.get("highlights")),
        "photo_data_url": str(raw.get("photo_data_url") or raw.get("photo_url") or "").strip(),
    }


def decode_photo(data_url: str) -> bytes | None:
    raw = (data_url or "").strip()
    if not raw:
        return None
    if raw.startswith("data:") and "," in raw:
        raw = raw.split(",", 1)[1]
    try:
        blob = base64.b64decode(raw, validate=False)
    except Exception:
        return None
    if not blob or len(blob) > _PHOTO_MAX:
        return None
    if blob[:3] not in (b"\xff\xd8\xff", b"\x89PN") and blob[:8] != b"\x89PNG\r\n\x1a\n":
        if not (blob.startswith(b"\xff\xd8") or blob.startswith(b"\x89PNG")):
            return None
    return blob


def _esc(text: str) -> str:
    return html.escape(text or "")


def _contact_line(data: dict[str, Any]) -> str:
    parts = [data["city"], data["contact"], data["email"], data["github"]]
    return " · ".join(p for p in parts if p)


def _chips_html(items: list[str]) -> str:
    if not items:
        return "<p class='muted'>待填写</p>"
    return "<div class='chips'>" + "".join(f"<span>{_esc(s)}</span>" for s in items) + "</div>"


def _ul_html(items: list[str], empty: str = "待填写") -> str:
    if not items:
        return f"<p class='muted'>{_esc(empty)}</p>"
    return "<ul>" + "".join(f"<li>{_esc(s)}</li>" for s in items) + "</ul>"


def _photo_html(data: dict[str, Any], allow: bool) -> str:
    if not allow:
        return ""
    src = data["photo_data_url"]
    if src.startswith("data:image"):
        img_src = src
    elif src.startswith("http") or src.startswith("/"):
        img_src = src
    elif src and decode_photo(src):
        img_src = f"data:image/jpeg;base64,{src}"
    else:
        return "<div class='photo photo--empty'>证件照</div>"
    return f"<div class='photo'><img src='{_esc(img_src)}' alt='证件照'></div>"


_CSS = """
@page { size: A4; margin: 12mm; }
* { box-sizing: border-box; }
body { margin: 0; background: #f4f1ea; font-family: "Source Han Sans SC","Noto Sans SC","PingFang SC","Microsoft YaHei",sans-serif; color: #1c1917; }
.sheet { width: 190mm; min-height: 277mm; margin: 12px auto; background: #fff; }
.muted { color: #78716c; font-size: 11px; }
.chips { display: flex; flex-wrap: wrap; gap: 6px; }
.chips span { border: 1px solid currentColor; border-radius: 999px; padding: 2px 8px; font-size: 11px; }
h2 { font-size: 12px; letter-spacing: .18em; text-transform: uppercase; margin: 14px 0 6px; }
ul { margin: 0; padding-left: 16px; font-size: 12.5px; line-height: 1.55; }
.photo { width: 26mm; height: 34mm; overflow: hidden; border: 1px solid #d6d3d1; background: #f5f5f4; display: flex; align-items: center; justify-content: center; font-size: 11px; color: #a8a29e; flex-shrink: 0; }
.photo img { width: 100%; height: 100%; object-fit: cover; }
.cv-hero { display: flex; justify-content: space-between; gap: 16px; align-items: flex-start; }
.cv-hero h1 { margin: 0; font-size: 26px; letter-spacing: .04em; }
.intent { margin: 4px 0 0; font-size: 13px; font-weight: 600; }
.meta { margin: 6px 0 0; font-size: 11.5px; color: #57534e; }
.grid { display: grid; grid-template-columns: 32% 1fr; gap: 18px; }
/* editorial */
.cv--editorial { padding: 14mm 14mm 12mm; }
.cv--editorial .cv-hero { border-top: 3px solid #b8860b; padding-top: 10px; }
.cv--editorial h2 { color: #b8860b; border-bottom: 1px solid #f0e6c8; padding-bottom: 3px; }
.cv--editorial .chips span { color: #92400e; border-color: #e7d3a1; }
/* navy */
.cv--navy_rail { display: grid; grid-template-columns: 58mm 1fr; min-height: 277mm; }
.cv--navy_rail .rail { background: #1e3a5f; color: #e8eef6; padding: 14mm 8mm; }
.cv--navy_rail .rail h1 { font-size: 22px; margin: 10px 0 4px; color: #fff; }
.cv--navy_rail .rail .intent, .cv--navy_rail .rail .meta { color: #cbd5e1; }
.cv--navy_rail .rail h2 { color: #fbbf24; letter-spacing: .14em; }
.cv--navy_rail .rail .chips span { color: #e2e8f0; border-color: #64748b; }
.cv--navy_rail .rail .photo { width: 32mm; height: 42mm; margin: 0 auto 8px; border-color: #fbbf24; }
.cv--navy_rail .main { padding: 14mm 12mm; }
.cv--navy_rail .main h2 { color: #1e3a5f; border-left: 3px solid #b8860b; padding-left: 8px; }
/* folio */
.cv--folio { padding: 14mm; }
.cv--folio h1 { font-family: "Source Han Serif SC","Songti SC","SimSun",serif; font-size: 24px; }
.cv--folio h2 { font-family: "Source Han Serif SC","Songti SC","SimSun",serif; color: #4c1d95; letter-spacing: .12em; }
.cv--folio .cv-hero { border-bottom: 1px solid #ddd6fe; padding-bottom: 10px; }
/* ats */
.cv--ats_plain { padding: 14mm; }
.cv--ats_plain h1 { font-size: 22px; }
.cv--ats_plain h2 { color: #0f172a; letter-spacing: 0; text-transform: none; border-bottom: 1px solid #cbd5e1; }
.cv--ats_plain .chips span { border: none; padding: 0; }
@media print { body { background: #fff; } .sheet { margin: 0; box-shadow: none; } }
"""


def build_resume_html(fields: dict[str, Any] | None, template_id: str = "editorial") -> str:
    spec = get_resume_template(template_id) or get_resume_template("editorial")
    tid = (spec or {}).get("id") or "editorial"
    allow = bool((spec or {}).get("allow_photo", True))
    data = normalize_fields(fields)
    name = _esc(data["name"])
    intent = _esc(data["intent"] or "求职意向")
    meta = _esc(_contact_line(data) or "电话 · 邮箱 · 城市")
    photo = _photo_html(data, allow)
    edu = _ul_html(data["education"])
    exp = _ul_html(data["experience"])
    proj = _ul_html(data["projects"])
    skills = _chips_html(data["skills"])
    certs = _ul_html(data["certificates"], "可填写 CET-6 / 专业证书")
    papers = _ul_html(data["papers"] or data["highlights"], "论文 / 竞赛")

    if tid == "navy_rail":
        inner = f"""
        <article class="cv cv--navy_rail">
          <aside class="rail">
            {photo}
            <h1>{name}</h1>
            <p class="intent">{intent}</p>
            <p class="meta">{meta}</p>
            <h2>技能</h2>{skills}
            <h2>证书</h2>{certs}
          </aside>
          <div class="main">
            <h2>教育</h2>{edu}
            <h2>实习经历</h2>{exp}
            <h2>项目经历</h2>{proj}
          </div>
        </article>"""
    elif tid == "ats_plain":
        inner = f"""
        <article class="cv cv--ats_plain">
          <h1>{name}</h1>
          <p class="intent">{intent}</p>
          <p class="meta">{meta}</p>
          <h2>教育背景</h2>{edu}
          <h2>实习经历</h2>{exp}
          <h2>项目经历</h2>{proj}
          <h2>技能</h2>{skills}
          <h2>证书</h2>{certs}
        </article>"""
    elif tid == "folio":
        inner = f"""
        <article class="cv cv--folio">
          <header class="cv-hero">
            <div>
              <h1>{name}</h1>
              <p class="intent">{intent}</p>
              <p class="meta">{meta}</p>
            </div>
            {photo}
          </header>
          <div class="grid">
            <aside>
              <h2>技能</h2>{skills}
              <h2>证书</h2>{certs}
            </aside>
            <main>
              <h2>教育</h2>{edu}
              <h2>科研 / 论文</h2>{papers}
              <h2>项目</h2>{proj}
              <h2>实习</h2>{exp}
            </main>
          </div>
        </article>"""
    else:
        inner = f"""
        <article class="cv cv--editorial">
          <header class="cv-hero">
            <div>
              <h1>{name}</h1>
              <p class="intent">{intent}</p>
              <p class="meta">{meta}</p>
            </div>
            {photo}
          </header>
          <div class="grid">
            <aside>
              <h2>技能</h2>{skills}
              <h2>证书</h2>{certs}
              <h2>城市</h2><p>{_esc(data["city"] or "—")}</p>
            </aside>
            <main>
              <h2>教育</h2>{edu}
              <h2>实习经历</h2>{exp}
              <h2>项目经历</h2>{proj}
            </main>
          </div>
        </article>"""

    title = f"{data['name']} · {(spec or {}).get('name') or '简历'}"
    return (
        "<!DOCTYPE html><html lang='zh-CN'><head><meta charset='utf-8'>"
        f"<title>{_esc(title)}</title><style>{_CSS}</style></head>"
        f"<body><div class='sheet'>{inner}</div></body></html>"
    )


def export_resume_md(fields: dict[str, Any] | None, template_id: str = "editorial") -> str:
    spec = get_resume_template(template_id) or get_resume_template("editorial")
    data = normalize_fields(fields)
    lines = [
        f"# {data['name']}",
        data["intent"] or "",
        _contact_line(data),
        "",
        "## 教育",
        *([f"- {x}" for x in data["education"]] or ["- "]),
        "",
        "## 实习",
        *([f"- {x}" for x in data["experience"]] or ["- "]),
        "",
        "## 项目",
        *([f"- {x}" for x in data["projects"]] or ["- "]),
        "",
        "## 技能",
        ", ".join(data["skills"]),
        "",
        "## 证书",
        *([f"- {x}" for x in data["certificates"]] or ["- "]),
        "",
        f"> 模板：{(spec or {}).get('name')}",
    ]
    return "\n".join(lines)


def export_resume_docx(fields: dict[str, Any] | None, template_id: str = "editorial") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    spec = get_resume_template(template_id) or get_resume_template("editorial")
    tid = resolve_template_id((spec or {}).get("id") or "editorial")
    allow = bool((spec or {}).get("allow_photo", True)) and tid != "ats_plain"
    data = normalize_fields(fields)
    photo = decode_photo(data["photo_data_url"]) if allow else None

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)

    accent = RGBColor(0x1E, 0x3A, 0x5F) if tid == "navy_rail" else RGBColor(0xB8, 0x86, 0x0B)
    if tid == "folio":
        accent = RGBColor(0x4C, 0x1D, 0x95)
    if tid == "ats_plain":
        accent = RGBColor(0x0F, 0x17, 0x2A)

    def _run_font(run, *, size=11, bold=False, color=None, name="Calibri"):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if color is not None:
            run.font.color.rgb = color

    def _heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _run_font(run, size=11, bold=True, color=accent)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    def _bullets(items: list[str]) -> None:
        if not items:
            p = doc.add_paragraph("待填写")
            if p.runs:
                _run_font(p.runs[0], size=10, color=RGBColor(0x78, 0x71, 0x6C))
            return
        for item in items:
            p = doc.add_paragraph(item, style="List Bullet")
            for run in p.runs:
                _run_font(run, size=10.5)

    def _fill_sidebar(cell, *, with_photo: bool, with_identity: bool) -> None:
        if with_photo and photo:
            try:
                cell.paragraphs[0].add_run().add_picture(BytesIO(photo), width=Cm(3.0))
            except Exception:
                pass
        if with_identity:
            n = cell.add_paragraph(data["name"])
            _run_font(n.runs[0], size=16, bold=True, color=accent)
            i = cell.add_paragraph(data["intent"] or "求职意向")
            _run_font(i.runs[0], size=10, bold=True)
            c = cell.add_paragraph(_contact_line(data))
            if c.runs:
                _run_font(c.runs[0], size=9)
        hs = cell.add_paragraph("技能")
        _run_font(hs.runs[0], size=10, bold=True, color=accent)
        for sk in data["skills"] or ["待填写"]:
            sp = cell.add_paragraph(sk, style="List Bullet")
            for run in sp.runs:
                _run_font(run, size=9)
        hc = cell.add_paragraph("证书")
        _run_font(hc.runs[0], size=10, bold=True, color=accent)
        for sk in data["certificates"] or ["CET 等"]:
            sp = cell.add_paragraph(sk, style="List Bullet")
            for run in sp.runs:
                _run_font(run, size=9)

    def _fill_story(cell, *, show_name: bool) -> None:
        if show_name:
            rp = cell.paragraphs[0]
            run = rp.add_run(data["name"])
            _run_font(run, size=20, bold=True, color=accent)
        def _rhead(text: str) -> None:
            p = cell.add_paragraph(text)
            _run_font(p.runs[0], size=11, bold=True, color=accent)

        def _rbullets(items: list[str]) -> None:
            for item in items or ["待填写"]:
                p = cell.add_paragraph(item, style="List Bullet")
                for r in p.runs:
                    _run_font(r, size=10.5)

        if tid == "folio":
            _rhead("教育")
            _rbullets(data["education"])
            _rhead("科研 / 论文")
            _rbullets(data["papers"] or data["highlights"])
            _rhead("项目")
            _rbullets(data["projects"])
            _rhead("实习")
            _rbullets(data["experience"])
        else:
            _rhead("教育")
            _rbullets(data["education"])
            _rhead("实习经历")
            _rbullets(data["experience"])
            _rhead("项目经历")
            _rbullets(data["projects"])

    if tid == "ats_plain":
        p = doc.add_paragraph()
        run = p.add_run(data["name"])
        _run_font(run, size=22, bold=True)
        p2 = doc.add_paragraph(data["intent"] or "求职意向")
        _run_font(p2.runs[0], size=11, bold=True)
        p3 = doc.add_paragraph(_contact_line(data))
        if p3.runs:
            _run_font(p3.runs[0], size=10)
        _heading("教育背景")
        _bullets(data["education"])
        _heading("实习经历")
        _bullets(data["experience"])
        _heading("项目经历")
        _bullets(data["projects"])
        _heading("技能")
        _bullets(data["skills"])
        _heading("证书")
        _bullets(data["certificates"])
    elif tid == "navy_rail":
        table = doc.add_table(rows=1, cols=2)
        left, right = table.rows[0].cells
        left.width = Cm(5.6)
        right.width = Cm(12.2)
        _fill_sidebar(left, with_photo=True, with_identity=True)
        _fill_story(right, show_name=False)
    else:
        header = doc.add_table(rows=1, cols=2)
        h_left, h_right = header.rows[0].cells
        h_left.width = Cm(13.8)
        h_right.width = Cm(4.0)
        hp = h_left.paragraphs[0]
        run = hp.add_run(data["name"])
        _run_font(run, size=22, bold=True, color=accent)
        i = h_left.add_paragraph(data["intent"] or "求职意向")
        _run_font(i.runs[0], size=11, bold=True)
        c = h_left.add_paragraph(_contact_line(data))
        if c.runs:
            _run_font(c.runs[0], size=10)
        if photo:
            rp = h_right.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                rp.add_run().add_picture(BytesIO(photo), width=Cm(2.8))
            except Exception:
                pass
        body = doc.add_table(rows=1, cols=2)
        left, right = body.rows[0].cells
        left.width = Cm(5.6)
        right.width = Cm(12.2)
        _fill_sidebar(left, with_photo=False, with_identity=False)
        _fill_story(right, show_name=False)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_resume(fields: dict[str, Any] | None, template_id: str, fmt: str) -> tuple[bytes, str, str]:
    """返回 (blob, media_type, filename)."""
    data = normalize_fields(fields)
    slug = re.sub(r"[^\w\u4e00-\u9fff-]+", "", data["name"]) or "resume"
    tid = resolve_template_id(template_id)
    kind = (fmt or "html").lower()
    if kind == "md":
        text = export_resume_md(fields, tid)
        return text.encode("utf-8"), "text/markdown; charset=utf-8", f"{slug}.md"
    if kind == "docx":
        blob = export_resume_docx(fields, tid)
        return (
            blob,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{slug}.docx",
        )
    html_doc = build_resume_html(fields, tid)
    return html_doc.encode("utf-8"), "text/html; charset=utf-8", f"{slug}.html"
