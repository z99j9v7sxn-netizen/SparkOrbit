"""将 作品设计实现方案.md 转换为 Word (.docx)，并把 Mermaid 图替换为已渲染 PNG。

用法：
    python build_docx.py
输出：
    docs/word/作品设计实现方案_多模型协同版.docx

说明：修改根目录 docs/作品设计实现方案.md（含「已实现/规划中」对照）后请重跑本脚本，
使 Word 交付物与代码口径一致，避免现场文档打脸。
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parents[1]
SRC_MD = DOCS_DIR / "作品设计实现方案.md"
IMG_DIR = Path(__file__).resolve().parent / "img"
OUT_DOCX = Path(__file__).resolve().parent / "作品设计实现方案_多模型协同版.docx"
# 若原文件被占用（如 Word 打开），回退到带时间戳的新文件名
OUT_DOCX_FALLBACK = Path(__file__).resolve().parent / "作品设计实现方案_多模型协同版_new.docx"

CN_FONT = "微软雅黑"
CODE_FONT = "Consolas"
MAX_IMG_WIDTH_CM = 15.5


def set_cn_font(run, name: str = CN_FONT, size: int | None = None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _ensure_rpr(run):
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn("w:rFonts")) is None:
        from docx.oxml import OxmlElement

        rpr.append(OxmlElement("w:rFonts"))


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def clean_inline(text: str) -> str:
    # markdown 链接 → 纯文本
    return LINK_RE.sub(r"\1", text)


def add_runs(paragraph, text: str, base_size: int | None = None, base_bold=False):
    text = clean_inline(text)
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            _ensure_rpr(r)
            set_cn_font(r, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            _ensure_rpr(r)
            set_cn_font(r, name=CODE_FONT, size=base_size)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            r = paragraph.add_run(part)
            _ensure_rpr(r)
            set_cn_font(r, size=base_size, bold=base_bold)


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {0: 22, 1: 17, 2: 14, 3: 12}
    r = p.add_run(clean_inline(text))
    _ensure_rpr(r)
    set_cn_font(r, size=sizes.get(level, 12), bold=True,
                color=RGBColor(0x1F, 0x3A, 0x5F) if level < 3 else None)
    p.space_after = Pt(6)
    return p


def add_image(doc, img_path: Path):
    if not img_path.exists():
        doc.add_paragraph(f"[缺失图片: {img_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    from PIL import Image

    with Image.open(img_path) as im:
        w, h = im.size
    width_cm = min(MAX_IMG_WIDTH_CM, w / 96 * 2.54)
    run.add_picture(str(img_path), width=Cm(width_cm))


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, val, base_size=9, base_bold=(i == 0))


def parse_table_block(lines: list[str], start: int):
    block = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        block.append(lines[i].strip())
        i += 1
    rows = []
    for idx, raw in enumerate(block):
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if idx == 1 and all(set(c) <= set("-: ") for c in cells):
            continue  # 分隔行
        rows.append(cells)
    return rows, i


def main():
    md = SRC_MD.read_text(encoding="utf-8")
    lines = md.split("\n")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    mermaid_idx = 0
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 分隔线
        if stripped == "---":
            i += 1
            continue

        # 代码围栏
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            j = i + 1
            code_lines = []
            while j < n and not lines[j].strip().startswith("```"):
                code_lines.append(lines[j])
                j += 1
            if lang == "mermaid":
                mermaid_idx += 1
                add_image(doc, IMG_DIR / f"diagram-{mermaid_idx}.png")
            else:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_lines))
                _ensure_rpr(r)
                set_cn_font(r, name=CODE_FONT, size=9)
            i = j + 1
            continue

        # 标题
        if stripped.startswith("#"):
            m = re.match(r"(#+)\s+(.*)", stripped)
            if m:
                level = len(m.group(1)) - 1
                add_heading(doc, m.group(2), level)
                i += 1
                continue

        # 表格
        if stripped.startswith("|"):
            rows, ni = parse_table_block(lines, i)
            add_table(doc, rows)
            i = ni
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, stripped.lstrip("> ").strip(), base_size=10)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # 无序列表
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_runs(p, stripped)
        i += 1

    try:
        doc.save(str(OUT_DOCX))
        saved = OUT_DOCX
    except PermissionError:
        doc.save(str(OUT_DOCX_FALLBACK))
        saved = OUT_DOCX_FALLBACK
        print(f"原文件被占用，已另存为: {saved}")
    print(f"saved: {saved}  (mermaid images used: {mermaid_idx})")


if __name__ == "__main__":
    main()
