"""导出 参加挑战赛理由及作品亮点.md → .docx"""
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
import re

SRC = Path(__file__).resolve().parents[1] / "参加挑战赛理由及作品亮点.md"
OUT = Path(r"c:\Users\咸\Desktop\文档\参加挑战赛理由及作品亮点.docx")
FALLBACK = Path(r"c:\Users\咸\Desktop\文档\参加挑战赛理由及作品亮点_new.docx")

CN_FONT = "微软雅黑"
CODE_FONT = "Consolas"


def set_font(run, name=CN_FONT, size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def ensure_rpr(run):
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn("w:rFonts")) is None:
        rpr.append(OxmlElement("w:rFonts"))


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def clean(text):
    return LINK_RE.sub(r"\1", text)


def add_runs(para, text, base_size=None, base_bold=False):
    text = clean(text)
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            ensure_rpr(r)
            set_font(r, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            ensure_rpr(r)
            set_font(r, name=CODE_FONT, size=base_size)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            r = para.add_run(part)
            ensure_rpr(r)
            set_font(r, size=base_size, bold=base_bold)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {0: 22, 1: 17, 2: 14, 3: 12}
    r = p.add_run(clean(text))
    ensure_rpr(r)
    set_font(
        r,
        size=sizes.get(level, 12),
        bold=True,
        color=RGBColor(0x1F, 0x3A, 0x5F) if level < 3 else None,
    )
    p.space_after = Pt(6)


def add_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            add_runs(cell.paragraphs[0], val, base_size=9, base_bold=(i == 0))


def parse_table(lines, start):
    block = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        block.append(lines[i].strip())
        i += 1
    rows = []
    for idx, raw in enumerate(block):
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if idx == 1 and all(set(c) <= set("-: ") for c in cells):
            continue
        rows.append(cells)
    return rows, i


def main():
    md = SRC.read_text(encoding="utf-8")
    lines = md.split("\n")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        if s.startswith("#"):
            m = re.match(r"(#+)\s+(.*)", s)
            if m:
                level = len(m.group(1)) - 1
                add_heading(doc, m.group(2), level)
                i += 1
                continue
        if s.startswith("|"):
            rows, ni = parse_table(lines, i)
            add_table(doc, rows)
            i = ni
            continue
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, s.lstrip("> ").strip(), base_size=10)
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_runs(p, s)
        i += 1

    try:
        doc.save(str(OUT))
        saved = OUT
    except PermissionError:
        doc.save(str(FALLBACK))
        saved = FALLBACK
        print(f"原文件被占用，已另存为: {saved}")
    print(f"已保存: {saved}")


if __name__ == "__main__":
    main()
