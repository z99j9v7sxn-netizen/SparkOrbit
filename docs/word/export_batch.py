"""批量导出指定 Markdown → Word (.docx)"""
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
import re

DOCS_DIR = Path(__file__).resolve().parents[1] / "software-eng"
OUT_DIR = DOCS_DIR / "docx"

CN_FONT = "微软雅黑"
CODE_FONT = "Consolas"


def set_font(run, name=CN_FONT, size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def ensure_rpr(run):
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn("w:rFonts")) is None:
        rpr.append(OxmlElement("w:rFonts"))


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def clean(text):
    return LINK_RE.sub(r"\1", text)


def add_runs(para, text, base_size=None, base_bold=False):
    text = clean(text)
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            ensure_rpr(r)
            set_font(r, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            ensure_rpr(r)
            set_font(r, name=CODE_FONT, size=base_size)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            r = para.add_run(part)
            ensure_rpr(r)
            set_font(r, size=base_size, bold=base_bold)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {0: 22, 1: 17, 2: 14, 3: 12}
    r = p.add_run(clean(text))
    ensure_rpr(r)
    set_font(r, size=sizes.get(level, 12), bold=True, color=RGBColor(0x1F, 0x3A, 0x5F) if level < 3 else None)
    p.space_after = Pt(6)


def add_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            add_runs(cell.paragraphs[0], val, base_size=9, base_bold=(i == 0))


def parse_table(lines, start):
    block = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        block.append(lines[i].strip())
        i += 1
    rows = []
    for idx, raw in enumerate(block):
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if idx == 1 and all(set(c) <= set("-: ") for c in cells):
            continue
        rows.append(cells)
    return rows, i


def md_to_docx(src_path: Path, out_path: Path):
    md = src_path.read_text(encoding="utf-8")
    lines = md.split("\n")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        if s.startswith("```"):
            j = i + 1
            code_lines = []
            while j < n and not lines[j].strip().startswith("```"):
                code_lines.append(lines[j])
                j += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            ensure_rpr(r)
            set_font(r, name=CODE_FONT, size=9)
            i = j + 1
            continue
        if s.startswith("#"):
            m = re.match(r"(#+)\s+(.*)", s)
            if m:
                add_heading(doc, m.group(2), len(m.group(1)) - 1)
                i += 1
                continue
        if s.startswith("|"):
            rows, ni = parse_table(lines, i)
            add_table(doc, rows)
            i = ni
            continue
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, s.lstrip("> ").strip(), base_size=10)
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_runs(p, s)
        i += 1

    try:
        doc.save(str(out_path))
        print(f"OK: {src_path.name} -> {out_path.name} ({out_path.stat().st_size / 1024:.0f} KB)")
    except PermissionError:
        alt = out_path.with_stem(out_path.stem + "_new")
        doc.save(str(alt))
        print(f"BUSY: {src_path.name} -> {alt.name}")


def main():
    targets = [
        ("SparkOrbit-A1-可行性研究报告.md", "SparkOrbit-A1-可行性研究报告.docx"),
        ("SparkOrbit-A2-项目开发计划.md", "SparkOrbit-A2-项目开发计划.docx"),
        ("SparkOrbit-B1-软件需求说明书.md", "SparkOrbit-B1-软件需求说明书.docx"),
        ("SparkOrbit-B2-数据要求说明书.md", "SparkOrbit-B2-数据要求说明书.docx"),
        ("SparkOrbit-C1-概要设计说明书.md", "SparkOrbit-C1-概要设计说明书.docx"),
        ("SparkOrbit-C2-详细设计说明书.md", "SparkOrbit-C2-详细设计说明书.docx"),
        ("SparkOrbit-C3-数据库设计说明书.md", "SparkOrbit-C3-数据库设计说明书.docx"),
        ("SparkOrbit-D1-模块开发卷宗.md", "SparkOrbit-D1-模块开发卷宗.docx"),
        ("SparkOrbit-D2-用户手册.md", "SparkOrbit-D2-用户手册.docx"),
        ("SparkOrbit-D3-操作手册.md", "SparkOrbit-D3-操作手册.docx"),
        ("SparkOrbit-E1-测试计划.md", "SparkOrbit-E1-测试计划.docx"),
        ("SparkOrbit-E2-测试分析报告.md", "SparkOrbit-E2-测试分析报告.docx"),
        ("SparkOrbit-F1-开发进度月报.md", "SparkOrbit-F1-开发进度月报.docx"),
        ("SparkOrbit-F2-项目开发总结报告.md", "SparkOrbit-F2-项目开发总结报告.docx"),
        ("SparkOrbit-G1-软件质量保证计划.md", "SparkOrbit-G1-软件质量保证计划.docx"),
        ("SparkOrbit-G2-软件配置管理计划.md", "SparkOrbit-G2-软件配置管理计划.docx"),
    ]
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for src_name, out_name in targets:
        src = DOCS_DIR / src_name
        if not src.exists():
            print(f"MISSING: {src}")
            continue
        out = OUT_DIR / out_name
        md_to_docx(src, out)


if __name__ == "__main__":
    main()
